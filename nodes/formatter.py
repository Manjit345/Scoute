"""
Formatter Node: Converts the synthesized research report into the user's chosen output format as PDF, DOCX, or Markdown and produces them in a downloadable file.
"""

import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from state.research_state import ResearchState

def format_as_markdown(report: str, output_dir: str, filename: str) -> str:
    """Save report as a plain markdown file."""
    path = os.path.join(output_dir, f"{filename}.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write(report)
    return path

def format_as_docx(report: str, output_dir: str, filename: str) -> str:
    """Save report as a word document along with converting markdown headers as docx headings"""
    doc = Document()

    for line in report.split('\n'):
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.strip():
            doc.add_paragraph(line)
       
    path = os.path.join(output_dir, f"{filename}.docx")
    doc.save(path)
    return path

def format_as_pdf(report: str, output_dir: str, filename: str) -> str:
    """Save report as a PDF document along with converting markdown headers into PDF styles"""
    path = os.path.join(output_dir, f"{filename}.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    for line in report.split('\n'):
        if line.startswith('### '):
            story.append(Paragraph(line.replace('### ', ''), styles['Heading3']))
        elif line.startswith('## '):
            story.append(Paragraph(line.replace('## ', ''), styles['Heading2']))
        elif line.startswith('# '):
            story.append(Paragraph(line.replace('# ', ''), styles['Heading1']))
        elif line.strip():
            story.append(Paragraph(line, styles['Normal']))
        story.append(Spacer(1, 6))
    
    doc.build(story)
    return path

def formatter(state: ResearchState) -> dict:
    """
    Convert the synthesized report into the requested output format.

    Args:
        state: Current research state containing the report and output_format.

    Returns:
        dict: Updated state fields with the output file path.
    """

    output_dir = "outputs"
    os.makedirs(output_dir, exist_ok=True)
    
    #Creating a filename from the topic
    filename = state["topic"][:50].replace(" ", "_").replace("/","_")
    
    try:
        report = state["report"]
        output_format = state["output_format"]
    
        if output_format == "markdown":
            path = format_as_markdown(report, output_dir, filename)
        elif output_format == "docx":
            path = format_as_docx(report, output_dir, filename)
        elif output_format == "pdf":
            path = format_as_pdf(report, output_dir, filename)
        else:
            path = format_as_markdown(report, output_dir, filename)
    
        return {"output_path": path}
    except Exception as e:
        print(f"Error in formatter: {e}")
        return {
            "output_path": None
        }

#Code for unit testing the function
if __name__ == "__main__":
    test_state: ResearchState = {
        "topic": "The impact of AI on journalism",
        "depth": "quick",
        "search_queries": None,
        "web_results": None,
        "academic_results": None,
        "evaluated_sources": None,
        "should_search_more": False,
        "report": "## The Impact of AI on Journalism\n\n### Executive Summary\nThis is a test report.\n\n### Conclusion\nTest conclusion.",
        "output_format": "pdf",
        "output_path": None
    }

    result = formatter(test_state)
    print(f"Output saved to: {result['output_path']}")
